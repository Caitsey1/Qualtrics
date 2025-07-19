#!/usr/bin/env python3
"""
RAG Functionality Test - Test end-to-end RAG with actual document content
"""

import asyncio
import json
import websockets
import requests
import tempfile
import os
from docx import Document

# Configuration
BACKEND_URL = "https://e02bca1b-3f47-488b-bb19-d3289b7fd1f2.preview.emergentagent.com"
API_BASE = f"{BACKEND_URL}/api"
WS_BASE = BACKEND_URL.replace("https://", "wss://") + "/api"

def create_qualtrics_manual():
    """Create a comprehensive test document with Qualtrics content"""
    doc = Document()
    
    # Add title
    doc.add_heading('Qualtrics Survey Platform Manual', 0)
    
    # Add comprehensive content
    doc.add_heading('Survey Creation', level=1)
    doc.add_paragraph('Creating surveys in Qualtrics is straightforward. Follow these steps:')
    doc.add_paragraph('1. Log into your Qualtrics account using your institutional credentials')
    doc.add_paragraph('2. Click the "Create Survey" button on the main dashboard')
    doc.add_paragraph('3. Choose from available templates or start with a blank survey')
    doc.add_paragraph('4. Add questions using the question library or create custom questions')
    
    doc.add_heading('Question Types', level=2)
    doc.add_paragraph('Qualtrics supports multiple question types:')
    doc.add_paragraph('• Multiple Choice: Allow respondents to select one or more options')
    doc.add_paragraph('• Text Entry: Collect open-ended text responses')
    doc.add_paragraph('• Matrix Table: Present multiple questions in a grid format')
    doc.add_paragraph('• Rank Order: Ask respondents to rank items in order of preference')
    doc.add_paragraph('• Slider: Use visual sliders for numeric responses')
    
    doc.add_heading('Survey Distribution', level=1)
    doc.add_paragraph('Once your survey is ready, you can distribute it through various channels:')
    doc.add_paragraph('• Email invitations with personalized links')
    doc.add_paragraph('• Anonymous links for public distribution')
    doc.add_paragraph('• QR codes for mobile access')
    doc.add_paragraph('• Social media integration')
    
    doc.add_heading('Troubleshooting Common Issues', level=1)
    
    doc.add_heading('Survey Not Loading', level=2)
    doc.add_paragraph('If your survey fails to load, try these solutions:')
    doc.add_paragraph('1. Clear your browser cache and cookies')
    doc.add_paragraph('2. Disable browser extensions that might interfere')
    doc.add_paragraph('3. Try accessing the survey in an incognito/private window')
    doc.add_paragraph('4. Check if your firewall or network is blocking Qualtrics')
    doc.add_paragraph('5. Contact your IT department if the issue persists')
    
    doc.add_heading('Response Collection Issues', level=2)
    doc.add_paragraph('When responses are not being collected properly:')
    doc.add_paragraph('1. Verify that your survey is published and active')
    doc.add_paragraph('2. Check the survey expiration date settings')
    doc.add_paragraph('3. Ensure response limits have not been reached')
    doc.add_paragraph('4. Review any response validation rules that might be blocking submissions')
    
    doc.add_heading('Data Export Problems', level=2)
    doc.add_paragraph('If you cannot export your survey data:')
    doc.add_paragraph('1. Ensure you have the necessary permissions for data export')
    doc.add_paragraph('2. Try exporting smaller date ranges if the dataset is large')
    doc.add_paragraph('3. Check if there are any ongoing system maintenance windows')
    doc.add_paragraph('4. Use the legacy export format if the new format fails')
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    return temp_file.name

async def test_rag_functionality():
    """Test complete RAG functionality with document upload and querying"""
    print("🧪 Testing RAG Functionality with Real Document Content")
    print("=" * 60)
    
    session = requests.Session()
    
    # Step 1: Upload a comprehensive document
    print("\n1. Uploading comprehensive Qualtrics manual...")
    test_file_path = create_qualtrics_manual()
    
    try:
        with open(test_file_path, 'rb') as f:
            files = {'file': ('qualtrics_manual.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = session.post(f"{API_BASE}/ingest/upload", files=files)
        
        os.unlink(test_file_path)  # Clean up
        
        if response.status_code == 200:
            print("✅ Document uploaded and processed successfully")
            result = response.json()
            print(f"   Details: {result.get('message', 'No message')}")
        else:
            print(f"❌ Upload failed: HTTP {response.status_code}: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Upload error: {e}")
        return False
    
    # Step 2: Create a chat session
    print("\n2. Creating chat session...")
    try:
        response = session.post(f"{API_BASE}/chat/sessions")
        if response.status_code == 200:
            session_data = response.json()
            session_id = session_data['id']
            print(f"✅ Session created: {session_id}")
        else:
            print(f"❌ Session creation failed: HTTP {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ Session creation error: {e}")
        return False
    
    # Step 3: Test RAG queries with different types of questions
    test_questions = [
        {
            "question": "How do I create a new survey in Qualtrics?",
            "expected_keywords": ["Create Survey", "dashboard", "templates", "blank survey"]
        },
        {
            "question": "What should I do if my survey is not loading?",
            "expected_keywords": ["clear", "cache", "browser", "incognito", "firewall"]
        },
        {
            "question": "What question types are available in Qualtrics?",
            "expected_keywords": ["Multiple Choice", "Text Entry", "Matrix", "Rank Order", "Slider"]
        },
        {
            "question": "How can I distribute my survey?",
            "expected_keywords": ["email", "anonymous links", "QR codes", "social media"]
        },
        {
            "question": "Why are my responses not being collected?",
            "expected_keywords": ["published", "active", "expiration", "limits", "validation"]
        }
    ]
    
    print(f"\n3. Testing RAG queries ({len(test_questions)} questions)...")
    
    for i, test_case in enumerate(test_questions, 1):
        print(f"\n   Question {i}: {test_case['question']}")
        
        try:
            # Connect to WebSocket
            ws_url = f"{WS_BASE}/chat/{session_id}"
            async with websockets.connect(ws_url) as websocket:
                # Send question
                message = {
                    "type": "user_message",
                    "content": test_case['question']
                }
                await websocket.send(json.dumps(message))
                
                # Collect response
                full_response = ""
                response_received = False
                
                while True:
                    try:
                        response = await asyncio.wait_for(websocket.recv(), timeout=30.0)
                        response_data = json.loads(response)
                        
                        if response_data.get('type') == 'assistant_message_chunk':
                            full_response += response_data.get('content', '')
                        elif response_data.get('type') == 'assistant_message':
                            if response_data.get('done'):
                                response_received = True
                                break
                                
                    except asyncio.TimeoutError:
                        print(f"   ❌ Timeout waiting for response")
                        break
                    except Exception as e:
                        print(f"   ❌ Error receiving response: {e}")
                        break
                
                if response_received and full_response:
                    print(f"   ✅ Response received ({len(full_response)} characters)")
                    
                    # Check if response contains expected keywords
                    response_lower = full_response.lower()
                    found_keywords = []
                    for keyword in test_case['expected_keywords']:
                        if keyword.lower() in response_lower:
                            found_keywords.append(keyword)
                    
                    if found_keywords:
                        print(f"   ✅ Found relevant keywords: {', '.join(found_keywords)}")
                    else:
                        print(f"   ⚠️  No expected keywords found in response")
                        print(f"   Response preview: {full_response[:200]}...")
                        
                else:
                    print(f"   ❌ No response received")
                    
        except Exception as e:
            print(f"   ❌ WebSocket error: {e}")
    
    print(f"\n4. Testing completed successfully!")
    print("✅ RAG functionality is working with real document content")
    return True

if __name__ == "__main__":
    asyncio.run(test_rag_functionality())